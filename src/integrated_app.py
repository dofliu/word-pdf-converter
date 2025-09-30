#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word與PDF轉換與合併應用程式 - 整合版
功能：
1. 將Word文件轉換為PDF格式
2. 將PDF文件轉換為Word格式
3. 將多個Word或PDF文件合併為單一PDF檔案
"""

import os
import sys
import time
import threading
import tempfile
import shutil
import platform
import subprocess
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QLabel, 
                            QFileDialog, QProgressBar, QMessageBox, QVBoxLayout, 
                            QHBoxLayout, QWidget, QGroupBox, QListWidget, QCheckBox,
                            QComboBox, QSpinBox, QListWidgetItem, QAbstractItemView,
                            QTabWidget, QTextEdit, QInputDialog, QLineEdit)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QUrl, QMetaObject, pyqtSlot, Q_ARG, Q_RETURN_ARG
from PyQt5.QtGui import QFont, QDesktopServices, QIcon
import docx2pdf
from docx import Document
import pypdf
from pypdf.errors import FileNotDecryptedError, WrongPasswordError
from pdf2docx import Converter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import cm  # 明確導入 cm 單位

# --- 常數 ---
USER_CANCELLED_PASSWORD = "USER_CANCELLED_MAGIC_STRING"

# 根據不同作業系統設定字型
def setup_fonts():
    """設定繁體中文字型，根據不同作業系統選擇適當的字型"""
    system = platform.system()
    
    if system == 'Windows':
        # Windows系統使用內建的微軟正黑體
        try:
            font_path = "C:/Windows/Fonts/msjh.ttc"
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                return 'ChineseFont'
            else:
                # 備用字型
                font_path = "C:/Windows/Fonts/simsun.ttc"
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    return 'ChineseFont'
                else:
                    # 如果找不到中文字型，使用預設字型
                    return 'Helvetica'
        except Exception as e:
            print(f"字型載入錯誤: {str(e)}")
            return 'Helvetica'
    
    elif system == 'Linux':
        # Linux系統嘗試使用Noto Sans CJK或文泉驛正黑
        try:
            font_paths = [
                '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                '/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc'
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    return 'ChineseFont'
            
            # 如果找不到中文字型，使用預設字型
            return 'Helvetica'
        except Exception as e:
            print(f"字型載入錯誤: {str(e)}")
            return 'Helvetica'
    
    elif system == 'Darwin':  # macOS
        # macOS系統嘗試使用蘋方或黑體
        try:
            font_paths = [
                '/System/Library/Fonts/PingFang.ttc',
                '/System/Library/Fonts/STHeiti Light.ttc',
                '/Library/Fonts/Arial Unicode.ttf'
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    return 'ChineseFont'
            
            # 如果找不到中文字型，使用預設字型
            return 'Helvetica'
        except Exception as e:
            print(f"字型載入錯誤: {str(e)}")
            return 'Helvetica'
    
    else:
        # 其他作業系統使用預設字型
        return 'Helvetica'

# 設定字型
CHINESE_FONT = setup_fonts()

# 增強版 Word 轉 PDF 函數
def convert_word_to_pdf(word_path, pdf_path):
    """使用多種方法嘗試將Word轉換為PDF，提高成功率"""
    system = platform.system()
    
    # 確保路徑是絕對路徑
    word_path = os.path.abspath(word_path)
    pdf_path = os.path.abspath(pdf_path)
    
    print(f"開始轉換: {word_path} -> {pdf_path}")
    
    # 方法1: 使用docx2pdf
    try:
        print("嘗試使用docx2pdf轉換...")
        docx2pdf.convert(word_path, pdf_path)
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            print("docx2pdf轉換成功!")
            return True
    except Exception as e:
        print(f"docx2pdf轉換失敗: {str(e)}")
    
    # 方法2: 使用COM自動化 (Windows專用)
    if system == 'Windows':
        try:
            print("嘗試使用COM自動化轉換...")
            import win32com.client
            import pythoncom
            
            # 初始化COM
            pythoncom.CoInitialize()
            
            # 創建Word應用實例
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                # 打開文檔
                doc = word.Documents.Open(word_path)
                # 另存為PDF (17代表PDF格式)
                doc.SaveAs(pdf_path, FileFormat=17)
                doc.Close()
                print("COM自動化轉換成功!")
                return True
            except Exception as e:
                print(f"Word COM轉換過程失敗: {str(e)}")
                return False
            finally:
                # 確保Word應用關閉
                try:
                    word.Quit()
                except:
                    pass
                # 釋放COM資源
                pythoncom.CoUninitialize()
        except Exception as e:
            print(f"Word COM初始化失敗: {str(e)}")
    
    # 方法3: 使用LibreOffice (如果安裝了)
    try:
        print("嘗試使用LibreOffice轉換...")
        # 檢查LibreOffice是否存在
        if system == 'Windows':
            soffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
            soffice_path = None
            for path in soffice_paths:
                if os.path.exists(path):
                    soffice_path = path
                    break
                    
            if soffice_path:
                cmd = [soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', 
                       os.path.dirname(pdf_path), word_path]
                process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                stdout, stderr = process.communicate()
                
                # 檢查轉換後的文件名稱
                base_name = os.path.splitext(os.path.basename(word_path))[0]
                temp_pdf = os.path.join(os.path.dirname(pdf_path), f"{base_name}.pdf")
                
                if os.path.exists(temp_pdf):
                    # 如果輸出路徑不同，移動文件
                    if temp_pdf != pdf_path:
                        shutil.move(temp_pdf, pdf_path)
                    print("LibreOffice轉換成功!")
                    return True
        else:
            # Linux/Mac檢查
            try:
                process = subprocess.Popen(['which', 'soffice'], stdout=subprocess.PIPE)
                soffice_path = process.communicate()[0].strip().decode('utf-8')
                
                if soffice_path:
                    cmd = [soffice_path, '--headless', '--convert-to', 'pdf', '--outdir', 
                           os.path.dirname(pdf_path), word_path]
                    process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    stdout, stderr = process.communicate()
                    
                    # 檢查轉換後的文件名稱
                    base_name = os.path.splitext(os.path.basename(word_path))[0]
                    temp_pdf = os.path.join(os.path.dirname(pdf_path), f"{base_name}.pdf")
                    
                    if os.path.exists(temp_pdf):
                        # 如果輸出路徑不同，移動文件
                        if temp_pdf != pdf_path:
                            shutil.move(temp_pdf, pdf_path)
                        print("LibreOffice轉換成功!")
                        return True
            except:
                pass
    except Exception as e:
        print(f"LibreOffice轉換失敗: {str(e)}")
    
    print("所有轉換方法都失敗了")
    return False


def convert_pdf_to_word(pdf_path, word_path):
    """將PDF轉換為Word文件"""
    try:
        print(f"開始轉換: {pdf_path} -> {word_path}")
        
        # 使用pdf2docx進行轉換
        cv = Converter(pdf_path)
        cv.convert(word_path, start=0, end=None)
        cv.close()
        
        # 檢查轉換結果
        if os.path.exists(word_path) and os.path.getsize(word_path) > 0:
            print("PDF轉Word成功!")
            return True
        else:
            print("PDF轉Word失敗: 輸出檔案為空")
            return False
    except Exception as e:
        print(f"PDF轉Word失敗: {str(e)}")
        return False


class PdfToWordThread(QThread):
    """處理PDF轉Word的執行緒，避免UI凍結"""
    progress_signal = pyqtSignal(int)
    status_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)
    
    def __init__(self, input_file, output_file):
        super().__init__()
        self.input_file = input_file
        self.output_file = output_file
        
    def run(self):
        try:
            # 更新狀態
            self.status_signal.emit("準備轉換...")
            self.progress_signal.emit(10)
            
            # 開始轉換
            self.status_signal.emit("轉換中...")
            self.progress_signal.emit(30)
            
            # 執行轉換
            success = convert_pdf_to_word(self.input_file, self.output_file)
            
            # 更新進度
            self.progress_signal.emit(90)
            
            # 檢查結果
            if success:
                self.status_signal.emit("轉換完成!")
                self.progress_signal.emit(100)
                self.finished_signal.emit(self.output_file)
            else:
                self.status_signal.emit("轉換失敗!")
                self.error_signal.emit("PDF轉Word失敗，請確認PDF檔案是否受保護或損壞。")
        except Exception as e:
            self.status_signal.emit("轉換出錯!")
            self.error_signal.emit(str(e))


class WordToPdfThread(QThread):
    """處理Word轉PDF的執行緒，避免UI凍結"""
    progress_signal = pyqtSignal(int)
    status_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)
    
    def __init__(self, input_file, output_file):
        super().__init__()
        self.input_file = input_file
        self.output_file = output_file
        
    def run(self):
        try:
            # 更新狀態
            self.status_signal.emit("準備轉換...")
            self.progress_signal.emit(10)
            
            # 開始轉換
            self.status_signal.emit("轉換中...")
            self.progress_signal.emit(30)
            
            # 執行轉換
            success = convert_word_to_pdf(self.input_file, self.output_file)
            
            # 更新進度
            self.progress_signal.emit(90)
            
            # 檢查結果
            if success and os.path.exists(self.output_file) and os.path.getsize(self.output_file) > 0:
                self.status_signal.emit("轉換完成!")
                self.progress_signal.emit(100)
                self.finished_signal.emit(self.output_file)
            else:
                self.status_signal.emit("轉換失敗!")
                self.error_signal.emit("轉換失敗，請確認Microsoft Word已正確安裝且可以開啟此檔案。")
        except Exception as e:
            self.status_signal.emit("轉換出錯!")
            self.error_signal.emit(str(e))


class MergePdfThread(QThread):
    """處理文件轉換與合併的執行緒，避免UI凍結"""
    progress_signal = pyqtSignal(int)
    status_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)
    
    def __init__(self, file_list, output_file, options, main_window):
        super().__init__()
        self.file_list = file_list
        self.output_file = output_file
        self.options = options
        # 确保主窗口对象引用是线程安全的
        self.main_window = main_window
        self.temp_dir = tempfile.mkdtemp()
        
    def _prepare_pdf_reader(self, pdf_path):
        """
        健壯地打開一個PDF文件。如果文件已加密，它將循環提示用戶輸入密碼，
        直到提供正確的密碼或用戶取消為止。
        """
        reader = pypdf.PdfReader(pdf_path)
        if not reader.is_encrypted:
            return reader

        # 文件已加密，需要解密
        filename = os.path.basename(pdf_path)
        while True:
            # 使用信号和槽机制安全地获取密码
            self.main_window.password_requested.emit(filename)
            password, ok = self.main_window.wait_for_password()
            
            if not ok:
                password = USER_CANCELLED_PASSWORD

            if password == USER_CANCELLED_PASSWORD:
                raise Exception(f"已取消操作，因為未提供 '{filename}' 的密碼。")

            try:
                if reader.decrypt(password):
                    return reader
                else:
                    self.main_window.wrong_password.emit(filename)
            except Exception as e:
                self.main_window.decrypt_error.emit(f"解密檔案 '{filename}' 時發生錯誤：{e}")
                raise

    def run(self):
        try:
            total_files = len(self.file_list)
            total_steps = total_files + 3  # 準備 + 目錄 + 合併 + 頁碼
            current_step = 0

            # 步驟 1: 準備階段 - 轉換Word並解密所有PDF
            self.status_signal.emit("準備檔案中...")
            prepared_pdfs = [] # 儲存 {'reader': PdfReader物件, 'title': str}
            
            for i, file_info in enumerate(self.file_list):
                file_path = file_info['path']
                file_name = os.path.basename(file_path)
                self.status_signal.emit(f"處理檔案 {i+1}/{total_files}: {file_name}")

                pdf_to_open = None
                if file_path.lower().endswith(('.docx', '.doc')):
                    temp_pdf = os.path.join(self.temp_dir, f"temp_{i}.pdf")
                    if not convert_word_to_pdf(file_path, temp_pdf):
                        raise Exception(f"無法轉換 Word 檔案: {file_name}。")
                    pdf_to_open = temp_pdf
                else:
                    pdf_to_open = file_path
                
                reader = self._prepare_pdf_reader(pdf_to_open)
                prepared_pdfs.append({'reader': reader, 'title': file_info['title']})
                
                current_step += 1
                self.progress_signal.emit(int(current_step * 100 / total_steps))

            # 步驟 2: 建立目錄 (如果需要)
            toc_pdf_path = None
            if self.options['generate_toc']:
                self.status_signal.emit("生成目錄...")
                toc_pdf_path = self.create_toc(prepared_pdfs)
            current_step += 1
            self.progress_signal.emit(int(current_step * 100 / total_steps))

            # 步驟 3: 合併所有PDF
            self.status_signal.emit("合併PDF檔案...")
            merger = pypdf.PdfWriter()
            if toc_pdf_path:
                merger.append(toc_pdf_path)
            for pdf_info in prepared_pdfs:
                merger.append(fileobj=pdf_info['reader'])

            merged_pdf_path = os.path.join(self.temp_dir, "merged.pdf")
            merger.write(merged_pdf_path)
            merger.close()
            current_step += 1
            self.progress_signal.emit(int(current_step * 100 / total_steps))

            # 步驟 4: 添加頁碼 (如果需要)
            final_pdf_path = merged_pdf_path
            if self.options['add_page_numbers']:
                self.status_signal.emit("添加頁碼...")
                final_pdf_path = self.add_page_numbers(merged_pdf_path)
            current_step += 1
            self.progress_signal.emit(int(current_step * 100 / total_steps))
            
            # 步驟 5: 完成
            shutil.copy2(final_pdf_path, self.output_file)
            self.progress_signal.emit(100)
            self.status_signal.emit("完成!")
            self.finished_signal.emit(self.output_file)
            
        except Exception as e:
            self.error_signal.emit(str(e))
        finally:
            try:
                shutil.rmtree(self.temp_dir)
            except:
                pass
    
    def create_toc(self, prepared_pdfs):
        """使用已解密的PdfReader物件創建目錄頁"""
        toc_pdf_path = os.path.join(self.temp_dir, "toc.pdf")
        c = canvas.Canvas(toc_pdf_path, pagesize=A4)
        c.setFont(CHINESE_FONT, 24)
        c.drawString(4*cm, 27*cm, "目錄")
        
        y_position = 25*cm
        c.setFont(CHINESE_FONT, 12)
        
        current_page_count = 1  # 目錄頁本身
        
        for i, pdf_info in enumerate(prepared_pdfs):
            title = pdf_info['title']
            if len(title) > 50:
                title = title[:47] + "..."
                
            c.drawString(2*cm, y_position, f"{i+1}. {title}")
            c.drawString(16*cm, y_position, f"{current_page_count + 1}")
            
            y_position -= 0.8*cm
            if y_position < 2*cm:
                c.showPage()
                c.setFont(CHINESE_FONT, 12)
                y_position = 27*cm
            
            current_page_count += len(pdf_info['reader'].pages)
        
        c.save()
        return toc_pdf_path
    
    def add_page_numbers(self, pdf_path):
        """添加頁碼到PDF"""
        output_path = os.path.join(self.temp_dir, "with_page_numbers.pdf")
        reader = pypdf.PdfReader(pdf_path)
        writer = pypdf.PdfWriter()
        
        page_format = self.options['page_number_format']
        start_number = self.options['start_page_number']
        
        for i, page in enumerate(reader.pages):
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)
            
            packet = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            c = canvas.Canvas(packet.name, pagesize=(page_width, page_height))
            c.setFont(CHINESE_FONT, 10)
            
            page_number = start_number + i
            page_text = str(page_number)
            if page_format == "羅馬數字":
                page_text = self.to_roman(page_number)
            
            c.drawString(page_width/2 - 10, 20, page_text)
            c.save()
            
            watermark = pypdf.PdfReader(packet.name)
            page.merge_page(watermark.pages[0])
            writer.add_page(page)
            
            packet.close()
            os.unlink(packet.name)
        
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return output_path
    
    def to_roman(self, num):
        """將數字轉換為羅馬數字"""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
        roman_num = ''
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syms[i]
                num -= val[i]
            i += 1
        return roman_num


class WordToPdfTab(QWidget):
    """Word轉PDF標籤頁"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
        
    def init_ui(self):
        """初始化使用者界面"""
        # 主佈局
        main_layout = QVBoxLayout()
        
        # 檔案選擇區域
        file_group = QGroupBox('檔案選擇')
        file_layout = QVBoxLayout()
        
        input_layout = QHBoxLayout()
        self.input_label = QLabel('Word檔案:')
        self.input_path = QTextEdit()
        self.input_path.setReadOnly(True)
        self.input_path.setMaximumHeight(60)
        self.browse_btn = QPushButton('瀏覽...')
        self.browse_btn.clicked.connect(self.browse_file)
        
        input_layout.addWidget(self.input_label)
        input_layout.addWidget(self.input_path)
        input_layout.addWidget(self.browse_btn)
        
        output_layout = QHBoxLayout()
        self.output_label = QLabel('儲存位置:')
        self.output_path = QTextEdit()
        self.output_path.setReadOnly(True)
        self.output_path.setMaximumHeight(60)
        self.save_btn = QPushButton('選擇...')
        self.save_btn.clicked.connect(self.save_file)
        
        output_layout.addWidget(self.output_label)
        output_layout.addWidget(self.output_path)
        output_layout.addWidget(self.save_btn)
        
        file_layout.addLayout(input_layout)
        file_layout.addLayout(output_layout)
        file_group.setLayout(file_layout)
        
        # 檔案資訊區域
        info_group = QGroupBox('檔案資訊')
        info_layout = QVBoxLayout()
        self.file_info = QTextEdit()
        self.file_info.setReadOnly(True)
        info_layout.addWidget(self.file_info)
        info_group.setLayout(info_layout)
        
        # 轉換選項區域
        options_group = QGroupBox('轉換選項')
        options_layout = QVBoxLayout()
        self.keep_format_cb = QCheckBox('保留格式、圖片與表格')
        self.keep_format_cb.setChecked(True)
        self.keep_format_cb.setEnabled(False)  # 預設啟用且不可更改
        options_layout.addWidget(self.keep_format_cb)
        options_group.setLayout(options_layout)
        
        # 進度條區域
        progress_layout = QVBoxLayout()
        self.status_label = QLabel('轉換進度:')
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.status_label)
        progress_layout.addWidget(self.progress_bar)
        
        # 操作按鈕區域
        button_layout = QHBoxLayout()
        self.convert_btn = QPushButton('開始轉換')
        self.convert_btn.clicked.connect(self.start_conversion)
        self.convert_btn.setEnabled(False)
        
        self.open_btn = QPushButton('開啟檔案')
        self.open_btn.clicked.connect(self.open_output_file)
        self.open_btn.setEnabled(False)
        
        button_layout.addWidget(self.convert_btn)
        button_layout.addWidget(self.open_btn)
        
        # 添加所有元件到主佈局
        main_layout.addWidget(file_group)
        main_layout.addWidget(info_group)
        main_layout.addWidget(options_group)
        main_layout.addLayout(progress_layout)
        main_layout.addLayout(button_layout)
        
        self.setLayout(main_layout)
        
        # 初始化變數
        self.input_file = ''
        self.output_file = ''
        
    def browse_file(self):
        """選擇Word檔案"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, '選擇Word檔案', '', 'Word檔案 (*.docx *.doc)'
        )
        
        if file_path:
            self.input_file = file_path
            self.input_path.setText(file_path)
            
            # 自動設定輸出檔案路徑
            dir_name = os.path.dirname(file_path)
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            self.output_file = os.path.join(dir_name, f"{base_name}.pdf")
            self.output_path.setText(self.output_file)
            
            # 顯示檔案資訊
            self.show_file_info(file_path)
            
            # 啟用轉換按鈕
            self.convert_btn.setEnabled(True)
    
    def save_file(self):
        """選擇PDF儲存位置"""
        if not self.input_file:
            QMessageBox.warning(self, '警告', '請先選擇Word檔案')
            return
            
        base_name = os.path.splitext(os.path.basename(self.input_file))[0]
        file_path, _ = QFileDialog.getSaveFileName(
            self, '儲存PDF檔案', f"{base_name}.pdf", 'PDF檔案 (*.pdf)'
        )
        
        if file_path:
            self.output_file = file_path
            self.output_path.setText(file_path)
    
    def show_file_info(self, file_path):
        """顯示Word檔案資訊"""
        try:
            doc = Document(file_path)
            
            # 計算頁數（近似值）
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            sections = len(doc.sections)
            
            # 計算圖片數量
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            file_size = os.path.getsize(file_path) / 1024  # KB
            
            info = f"檔案名稱: {os.path.basename(file_path)}\n"
            info += f"檔案大小: {file_size:.2f} KB\n"
            info += f"段落數量: {paragraphs}\n"
            info += f"表格數量: {tables}\n"
            info += f"圖片數量: {image_count}\n"
            info += f"章節數量: {sections}\n"
            
            self.file_info.setText(info)
        except Exception as e:
            self.file_info.setText(f"無法讀取檔案資訊: {str(e)}")
    
    def start_conversion(self):
        """開始轉換Word到PDF"""
        if not self.input_file or not self.output_file:
            QMessageBox.warning(self, '警告', '請選擇輸入和輸出檔案')
            return
        
        # 禁用按鈕，避免重複操作
        self.convert_btn.setEnabled(False)
        self.browse_btn.setEnabled(False)
        self.save_btn.setEnabled(False)
        
        # 開始轉換執行緒
        self.conversion_thread = WordToPdfThread(self.input_file, self.output_file)
        self.conversion_thread.progress_signal.connect(self.update_progress)
        self.conversion_thread.status_signal.connect(self.update_status)
        self.conversion_thread.finished_signal.connect(self.conversion_finished)
        self.conversion_thread.error_signal.connect(self.conversion_error)
        self.conversion_thread.start()
    
    def update_progress(self, value):
        """更新進度條"""
        self.progress_bar.setValue(value)
    
    def update_status(self, status):
        """更新狀態標籤"""
        self.status_label.setText(status)
    
    def conversion_finished(self, output_file):
        """轉換完成處理"""
        self.status_label.setText('轉換完成!')
        
        # 重新啟用按鈕
        self.convert_btn.setEnabled(True)
        self.browse_btn.setEnabled(True)
        self.save_btn.setEnabled(True)
        self.open_btn.setEnabled(True)
        
        QMessageBox.information(
            self, '完成', f'Word檔案已成功轉換為PDF!\n儲存於: {output_file}'
        )
    
    def conversion_error(self, error_msg):
        """轉換錯誤處理"""
        self.status_label.setText('轉換失敗!')
        self.progress_bar.setValue(0)
        
        # 重新啟用按鈕
        self.convert_btn.setEnabled(True)
        self.browse_btn.setEnabled(True)
        self.save_btn.setEnabled(True)
        
        QMessageBox.critical(self, '錯誤', f'轉換過程中發生錯誤:\n{error_msg}')
    
    def open_output_file(self):
        """開啟生成的PDF檔案"""
        if os.path.exists(self.output_file):
            QDesktopServices.openUrl(QUrl.fromLocalFile(self.output_file))
        else:
            QMessageBox.warning(self, '警告', '找不到輸出檔案')


class PdfMergerTab(QWidget):
    """多文件合併PDF標籤頁"""
    
    def __init__(self, parent=None, main_window=None):
        super().__init__(parent)
        self.main_window = main_window
        self.init_ui()
        
    def init_ui(self):
        """初始化使用者界面"""
        # 主佈局
        main_layout = QVBoxLayout()
        
        # 檔案列表區域
        file_group = QGroupBox('檔案列表')
        file_layout = QVBoxLayout()
        
        # 檔案列表
        self.file_list = QListWidget()
        self.file_list.setDragDropMode(QAbstractItemView.InternalMove)  # 允許拖放調整順序
        
        # 檔案操作按鈕
        file_buttons_layout = QHBoxLayout()
        self.add_file_btn = QPushButton('添加檔案')
        self.add_file_btn.clicked.connect(self.add_files)
        
        self.remove_file_btn = QPushButton('移除檔案')
        self.remove_file_btn.clicked.connect(self.remove_file)
        
        self.move_up_btn = QPushButton('上移')
        self.move_up_btn.clicked.connect(self.move_file_up)
        
        self.move_down_btn = QPushButton('下移')
        self.move_down_btn.clicked.connect(self.move_file_down)
        
        file_buttons_layout.addWidget(self.add_file_btn)
        file_buttons_layout.addWidget(self.remove_file_btn)
        file_buttons_layout.addWidget(self.move_up_btn)
        file_buttons_layout.addWidget(self.move_down_btn)
        
        file_layout.addWidget(self.file_list)
        file_layout.addLayout(file_buttons_layout)
        file_group.setLayout(file_layout)
        
        # 合併選項區域
        options_group = QGroupBox('合併選項')
        options_layout = QVBoxLayout()
        
        # 目錄選項
        toc_layout = QHBoxLayout()
        self.generate_toc_cb = QCheckBox('生成目錄')
        self.generate_toc_cb.setChecked(True)
        toc_layout.addWidget(self.generate_toc_cb)
        
        # 頁碼選項
        page_number_layout = QHBoxLayout()
        self.add_page_numbers_cb = QCheckBox('添加頁碼')
        self.add_page_numbers_cb.setChecked(True)
        self.add_page_numbers_cb.stateChanged.connect(self.toggle_page_options)
        
        self.page_format_label = QLabel('頁碼格式:')
        self.page_format_combo = QComboBox()
        self.page_format_combo.addItems(['數字', '羅馬數字'])
        
        self.start_page_label = QLabel('起始頁碼:')
        self.start_page_spin = QSpinBox()
        self.start_page_spin.setMinimum(1)
        self.start_page_spin.setValue(1)
        
        page_number_layout.addWidget(self.add_page_numbers_cb)
        page_number_layout.addWidget(self.page_format_label)
        page_number_layout.addWidget(self.page_format_combo)
        page_number_layout.addWidget(self.start_page_label)
        page_number_layout.addWidget(self.start_page_spin)
        
        options_layout.addLayout(toc_layout)
        options_layout.addLayout(page_number_layout)
        options_group.setLayout(options_layout)
        
        # 輸出選項區域
        output_group = QGroupBox('輸出設定')
        output_layout = QHBoxLayout()
        
        self.output_label = QLabel('輸出檔案:')
        self.output_path = QLabel('尚未選擇輸出檔案')
        self.output_btn = QPushButton('選擇...')
        self.output_btn.clicked.connect(self.select_output)
        
        output_layout.addWidget(self.output_label)
        output_layout.addWidget(self.output_path, 1)  # 1表示可伸展
        output_layout.addWidget(self.output_btn)
        
        output_group.setLayout(output_layout)
        
        # 進度區域
        progress_layout = QVBoxLayout()
        self.status_label = QLabel('準備就緒')
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        
        progress_layout.addWidget(self.status_label)
        progress_layout.addWidget(self.progress_bar)
        
        # 操作按鈕區域
        button_layout = QHBoxLayout()
        self.merge_btn = QPushButton('開始合併')
        self.merge_btn.clicked.connect(self.start_merge)
        self.merge_btn.setEnabled(False)
        
        self.open_btn = QPushButton('開啟檔案')
        self.open_btn.clicked.connect(self.open_output_file)
        self.open_btn.setEnabled(False)
        
        button_layout.addWidget(self.merge_btn)
        button_layout.addWidget(self.open_btn)
        
        # 添加所有元件到主佈局
        main_layout.addWidget(file_group, 3)  # 3表示佔比較大
        main_layout.addWidget(options_group, 1)
        main_layout.addWidget(output_group, 1)
        main_layout.addLayout(progress_layout)
        main_layout.addLayout(button_layout)
        
        self.setLayout(main_layout)
        
        # 初始化變數
        self.files = []
        self.output_file = ''
        
        # 初始化頁碼選項狀態
        self.toggle_page_options()
    
    def add_files(self):
        """添加檔案到列表"""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, '選擇檔案', '', 'Word檔案 (*.docx *.doc);;PDF檔案 (*.pdf);;所有檔案 (*.*)'
        )
        
        if file_paths:
            for file_path in file_paths:
                # 檢查檔案是否已在列表中
                if any(file['path'] == file_path for file in self.files):
                    continue
                
                # 獲取檔案名稱作為標題
                title = os.path.splitext(os.path.basename(file_path))[0]
                
                # 添加到內部列表
                self.files.append({
                    'path': file_path,
                    'title': title
                })
                
                # 添加到UI列表
                display_text = f"{title} ({os.path.basename(file_path)})"
                item = QListWidgetItem(display_text)
                self.file_list.addItem(item)
            
            # 如果有檔案且已選擇輸出路徑，啟用合併按鈕
            if self.files and self.output_file:
                self.merge_btn.setEnabled(True)
    
    def remove_file(self):
        """從列表中移除選中的檔案"""
        current_row = self.file_list.currentRow()
        if current_row >= 0:
            self.file_list.takeItem(current_row)
            self.files.pop(current_row)
            
            # 如果列表為空或未選擇輸出路徑，禁用合併按鈕
            if not self.files or not self.output_file:
                self.merge_btn.setEnabled(False)
    
    def move_file_up(self):
        """將選中的檔案上移"""
        current_row = self.file_list.currentRow()
        if current_row > 0:
            # 移動UI列表項目
            item = self.file_list.takeItem(current_row)
            self.file_list.insertItem(current_row - 1, item)
            self.file_list.setCurrentRow(current_row - 1)
            
            # 移動內部列表項目
            self.files.insert(current_row - 1, self.files.pop(current_row))
    
    def move_file_down(self):
        """將選中的檔案下移"""
        current_row = self.file_list.currentRow()
        if current_row < self.file_list.count() - 1:
            # 移動UI列表項目
            item = self.file_list.takeItem(current_row)
            self.file_list.insertItem(current_row + 1, item)
            self.file_list.setCurrentRow(current_row + 1)
            
            # 移動內部列表項目
            self.files.insert(current_row + 1, self.files.pop(current_row))
    
    def toggle_page_options(self):
        """切換頁碼選項的啟用狀態"""
        enabled = self.add_page_numbers_cb.isChecked()
        self.page_format_label.setEnabled(enabled)
        self.page_format_combo.setEnabled(enabled)
        self.start_page_label.setEnabled(enabled)
        self.start_page_spin.setEnabled(enabled)
    
    def select_output(self):
        """選擇輸出檔案路徑"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, '儲存合併PDF', '', 'PDF檔案 (*.pdf)'
        )
        
        if file_path:
            # 確保有.pdf副檔名
            if not file_path.lower().endswith('.pdf'):
                file_path += '.pdf'
                
            self.output_file = file_path
            self.output_path.setText(os.path.basename(file_path))
            
            # 如果有檔案，啟用合併按鈕
            if self.files:
                self.merge_btn.setEnabled(True)
    
    def start_merge(self):
        """開始合併PDF檔案"""
        if not self.files:
            QMessageBox.warning(self, '警告', '請先添加檔案')
            return
            
        if not self.output_file:
            QMessageBox.warning(self, '警告', '請選擇輸出檔案路徑')
            return
        
        # 獲取合併選項
        options = {
            'generate_toc': self.generate_toc_cb.isChecked(),
            'add_page_numbers': self.add_page_numbers_cb.isChecked(),
            'page_number_format': self.page_format_combo.currentText(),
            'start_page_number': self.start_page_spin.value()
        }
        
        # 禁用按鈕，避免重複操作
        self.merge_btn.setEnabled(False)
        self.add_file_btn.setEnabled(False)
        self.remove_file_btn.setEnabled(False)
        self.move_up_btn.setEnabled(False)
        self.move_down_btn.setEnabled(False)
        self.output_btn.setEnabled(False)
        
        # 開始合併執行緒，並明確傳遞主視窗的參考
        self.conversion_thread = MergePdfThread(self.files, self.output_file, options, self.main_window)
        self.conversion_thread.progress_signal.connect(self.update_progress)
        self.conversion_thread.status_signal.connect(self.update_status)
        self.conversion_thread.finished_signal.connect(self.merge_finished)
        self.conversion_thread.error_signal.connect(self.merge_error)
        self.conversion_thread.start()
    
    def update_progress(self, value):
        """更新進度條"""
        self.progress_bar.setValue(value)
    
    def update_status(self, status):
        """更新狀態標籤"""
        self.status_label.setText(status)
    
    def merge_finished(self, output_file):
        """合併完成處理"""
        # 重新啟用按鈕
        self.merge_btn.setEnabled(True)
        self.add_file_btn.setEnabled(True)
        self.remove_file_btn.setEnabled(True)
        self.move_up_btn.setEnabled(True)
        self.move_down_btn.setEnabled(True)
        self.output_btn.setEnabled(True)
        self.open_btn.setEnabled(True)
        
        QMessageBox.information(
            self, '完成', f'檔案已成功合併!\n儲存於: {output_file}'
        )
    
    def merge_error(self, error_msg):
        """合併錯誤處理"""
        self.status_label.setText('合併失敗!')
        self.progress_bar.setValue(0)
        
        # 重新啟用按鈕
        self.merge_btn.setEnabled(True)
        self.add_file_btn.setEnabled(True)
        self.remove_file_btn.setEnabled(True)
        self.move_up_btn.setEnabled(True)
        self.move_down_btn.setEnabled(True)
        self.output_btn.setEnabled(True)
        
        QMessageBox.critical(self, '錯誤', f'合併過程中發生錯誤:\n{error_msg}')
    
    def open_output_file(self):
        """開啟生成的PDF檔案"""
        if os.path.exists(self.output_file):
            QDesktopServices.openUrl(QUrl.fromLocalFile(self.output_file))
        else:
            QMessageBox.warning(self, '警告', '找不到輸出檔案')


class AboutTab(QWidget):
    """關於標籤頁"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
        
    def init_ui(self):
        """初始化使用者界面"""
        # 主佈局
        main_layout = QVBoxLayout()
        
        # 標題
        title_label = QLabel('Word與PDF轉換與合併工具')
        title_font = QFont()
        title_font.setFamily('Microsoft JhengHei')
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignCenter)
        
        # 版本
        version_label = QLabel('版本: 1.4.0 (穩定版)')
        version_label.setAlignment(Qt.AlignCenter)
        
        # 說明
        desc_text = QTextEdit()
        desc_text.setReadOnly(True)
        desc_text.setHtml("""
        <h2>功能說明</h2>
        <p>本應用程式提供兩個主要功能：</p>
        <ol>
            <li><b>Word轉PDF</b>：將Word文件轉換為PDF格式，保留原始格式、圖片與表格</li>
            <li><b>多文件合併PDF</b>：將多個Word或PDF文件合併為單一PDF檔案，支援調整順序、生成目錄與添加頁碼</li>
        </ol>
        
        <h2>使用須知</h2>
        <ul>
            <li>使用Word轉PDF功能需要安裝Microsoft Word</li>
            <li>合併加密的PDF檔案時，會提示輸入密碼</li>
            <li>合併大型文件可能需要較長時間</li>
            <li>支援繁體中文</li>
            <li>如果轉換失敗，請確認Word可以正常開啟該檔案</li>
            <li>建議以管理員身份執行本程式</li>
        </ul>
        
        <h2>系統需求</h2>
        <ul>
            <li>作業系統：Windows 10或更新版本</li>
            <li>Microsoft Word 2010或更新版本</li>
        </ul>
        
        <h2>常見問題</h2>
        <p><b>問：轉換Word檔案時出現錯誤</b></p>
        <p>答：請確認以下事項：</p>
        <ul>
            <li>Microsoft Word已正確安裝且可以開啟該檔案</li>
            <li>檔案未被其他程式鎖定</li>
            <li>嘗試以管理員身份執行本程式</li>
            <li>關閉所有開啟的Word視窗後再試</li>
        </ul>
        
        <p><b>問：合併PDF時出現錯誤</b></p>
        <p>答：請確認以下事項：</p>
        <ul>
            <li>所有Word檔案都可以正常開啟</li>
            <li>PDF檔案未被損壞或密碼正確</li>
            <li>有足夠的磁碟空間</li>
        </ul>
        """)
        
        # 版權資訊
        copyright_label = QLabel('© 2025 Word與PDF轉換工具')
        copyright_label.setAlignment(Qt.AlignCenter)
        
        # 添加所有元件到主佈局
        main_layout.addWidget(title_label)
        main_layout.addWidget(version_label)
        main_layout.addWidget(desc_text)
        main_layout.addWidget(copyright_label)
        
        self.setLayout(main_layout)


class IntegratedApp(QMainWindow):
    """整合版應用程式主視窗"""
    
    def __init__(self):
        super().__init__()
        self.init_ui()
        
    def init_ui(self):
        """初始化使用者界面"""
        self.setWindowTitle('Word與PDF轉換工具')
        self.setGeometry(300, 300, 800, 600)
        self.setMinimumSize(700, 500)
        
        # 設定字型以支援繁體中文
        self.font = QFont()
        if platform.system() == 'Windows':
            self.font.setFamily('Microsoft JhengHei')  # 微軟正黑體
        elif platform.system() == 'Darwin':  # macOS
            self.font.setFamily('PingFang TC')  # 蘋方
        else:
            self.font.setFamily('Noto Sans CJK TC')  # Noto Sans
        self.font.setPointSize(10)
        QApplication.setFont(self.font)
        
        # 創建標籤頁
        self.tabs = QTabWidget()
        
        # 添加Word轉PDF標籤頁
        self.word_to_pdf_tab = WordToPdfTab(parent=self)
        self.tabs.addTab(self.word_to_pdf_tab, 'Word轉PDF')
        
        # 添加多文件合併PDF標籤頁，並明確傳遞主視窗參考
        self.pdf_merger_tab = PdfMergerTab(parent=self, main_window=self)
        self.tabs.addTab(self.pdf_merger_tab, '多文件合併PDF')
        
        # 添加關於標籤頁
        self.about_tab = AboutTab(parent=self)
        self.tabs.addTab(self.about_tab, '關於')
        
        # 設定主視窗
        self.setCentralWidget(self.tabs)

    # @pyqtSlot(str, result=str)
    # def get_password(self, filename):
    #     """彈出對話框讓使用者輸入密碼，並返回結果"""
    #     print(f"get_password called for file: {filename}")  # 添加日誌記錄
    #     password, ok = QInputDialog.getText(
    #         self,
    #         '需要密碼',
    #         f'檔案 "{filename}" 已加密，請輸入密碼：',
    #         QLineEdit.Password
    #     )
    #     if ok:
    #         return password
    #     return USER_CANCELLED_PASSWORD

    # @pyqtSlot(str)
    # def show_wrong_password_warning(self, filename):
    #     """顯示密碼錯誤的警告框"""
    #     print(f"show_wrong_password_warning called for file: {filename}")  # 添加日誌記錄
    #     QMessageBox.warning(self, '密碼錯誤', f"檔案 '{filename}' 的密碼不正確，請重試。")


if __name__ == '__main__':
    app = QApplication(sys.argv)
    integrated_app = IntegratedApp()
    integrated_app.show()
    sys.exit(app.exec_())