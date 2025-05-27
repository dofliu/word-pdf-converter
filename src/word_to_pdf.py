#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word轉PDF應用程式
功能：將Word文件轉換為PDF格式，保留格式、圖片與表格，支援繁體中文
"""

import os
import sys
import time
import threading
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QLabel, 
                            QFileDialog, QProgressBar, QMessageBox, QVBoxLayout, 
                            QHBoxLayout, QWidget, QGroupBox, QTextEdit, QCheckBox)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QUrl
from PyQt5.QtGui import QFont, QDesktopServices
import docx2pdf
from docx import Document

class ConversionThread(QThread):
    """處理Word轉PDF的執行緒，避免UI凍結"""
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)
    
    def __init__(self, input_file, output_file):
        super().__init__()
        self.input_file = input_file
        self.output_file = output_file
        
    def run(self):
        try:
            # 模擬進度（實際轉換無法獲取進度）
            for i in range(1, 101):
                if i == 10:
                    # 開始轉換
                    self.convert_thread = threading.Thread(target=self.convert_file)
                    self.convert_thread.start()
                
                self.progress_signal.emit(i)
                time.sleep(0.05)
                
                # 等待轉換完成
                if i == 100:
                    if hasattr(self, 'convert_thread'):
                        self.convert_thread.join()
                    
                    if hasattr(self, 'conversion_error'):
                        self.error_signal.emit(self.conversion_error)
                    else:
                        self.finished_signal.emit(self.output_file)
        except Exception as e:
            self.error_signal.emit(str(e))
    
    def convert_file(self):
        try:
            docx2pdf.convert(self.input_file, self.output_file)
        except Exception as e:
            self.conversion_error = str(e)


class WordToPdfConverter(QMainWindow):
    """Word轉PDF應用程式主視窗"""
    
    def __init__(self):
        super().__init__()
        self.init_ui()
        
    def init_ui(self):
        """初始化使用者界面"""
        self.setWindowTitle('Word轉PDF轉換器')
        self.setGeometry(300, 300, 600, 400)
        self.setMinimumSize(500, 350)
        
        # 設定字型以支援繁體中文
        self.font = QFont()
        self.font.setFamily('Microsoft JhengHei')  # 微軟正黑體
        self.font.setPointSize(10)
        QApplication.setFont(self.font)
        
        # 主佈局
        main_layout = QVBoxLayout()
        
        # 檔案選擇區域
        file_group = QGroupBox('檔案選擇')
        file_layout = QVBoxLayout()
        
        input_layout = QHBoxLayout()
        self.input_label = QLabel('Word檔案:')
        self.input_path = QTextEdit()
        self.input_path.setReadOnly(True)
        self.input_path.setMaximumHeight(60)
        self.browse_btn = QPushButton('瀏覽...')
        self.browse_btn.clicked.connect(self.browse_file)
        
        input_layout.addWidget(self.input_label)
        input_layout.addWidget(self.input_path)
        input_layout.addWidget(self.browse_btn)
        
        output_layout = QHBoxLayout()
        self.output_label = QLabel('儲存位置:')
        self.output_path = QTextEdit()
        self.output_path.setReadOnly(True)
        self.output_path.setMaximumHeight(60)
        self.save_btn = QPushButton('選擇...')
        self.save_btn.clicked.connect(self.save_file)
        
        output_layout.addWidget(self.output_label)
        output_layout.addWidget(self.output_path)
        output_layout.addWidget(self.save_btn)
        
        file_layout.addLayout(input_layout)
        file_layout.addLayout(output_layout)
        file_group.setLayout(file_layout)
        
        # 檔案資訊區域
        info_group = QGroupBox('檔案資訊')
        info_layout = QVBoxLayout()
        self.file_info = QTextEdit()
        self.file_info.setReadOnly(True)
        info_layout.addWidget(self.file_info)
        info_group.setLayout(info_layout)
        
        # 轉換選項區域
        options_group = QGroupBox('轉換選項')
        options_layout = QVBoxLayout()
        self.keep_format_cb = QCheckBox('保留格式、圖片與表格')
        self.keep_format_cb.setChecked(True)
        self.keep_format_cb.setEnabled(False)  # 預設啟用且不可更改
        options_layout.addWidget(self.keep_format_cb)
        options_group.setLayout(options_layout)
        
        # 進度條區域
        progress_layout = QVBoxLayout()
        self.progress_label = QLabel('轉換進度:')
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.progress_label)
        progress_layout.addWidget(self.progress_bar)
        
        # 操作按鈕區域
        button_layout = QHBoxLayout()
        self.convert_btn = QPushButton('開始轉換')
        self.convert_btn.clicked.connect(self.start_conversion)
        self.convert_btn.setEnabled(False)
        
        self.open_btn = QPushButton('開啟檔案')
        self.open_btn.clicked.connect(self.open_output_file)
        self.open_btn.setEnabled(False)
        
        button_layout.addWidget(self.convert_btn)
        button_layout.addWidget(self.open_btn)
        
        # 添加所有元件到主佈局
        main_layout.addWidget(file_group)
        main_layout.addWidget(info_group)
        main_layout.addWidget(options_group)
        main_layout.addLayout(progress_layout)
        main_layout.addLayout(button_layout)
        
        # 設定主視窗
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)
        
        # 初始化變數
        self.input_file = ''
        self.output_file = ''
        
    def browse_file(self):
        """選擇Word檔案"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, '選擇Word檔案', '', 'Word檔案 (*.docx *.doc)'
        )
        
        if file_path:
            self.input_file = file_path
            self.input_path.setText(file_path)
            
            # 自動設定輸出檔案路徑
            dir_name = os.path.dirname(file_path)
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            self.output_file = os.path.join(dir_name, f"{base_name}.pdf")
            self.output_path.setText(self.output_file)
            
            # 顯示檔案資訊
            self.show_file_info(file_path)
            
            # 啟用轉換按鈕
            self.convert_btn.setEnabled(True)
    
    def save_file(self):
        """選擇PDF儲存位置"""
        if not self.input_file:
            QMessageBox.warning(self, '警告', '請先選擇Word檔案')
            return
            
        base_name = os.path.splitext(os.path.basename(self.input_file))[0]
        file_path, _ = QFileDialog.getSaveFileName(
            self, '儲存PDF檔案', f"{base_name}.pdf", 'PDF檔案 (*.pdf)'
        )
        
        if file_path:
            self.output_file = file_path
            self.output_path.setText(file_path)
    
    def show_file_info(self, file_path):
        """顯示Word檔案資訊"""
        try:
            doc = Document(file_path)
            
            # 計算頁數（近似值）
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            sections = len(doc.sections)
            
            # 計算圖片數量
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            file_size = os.path.getsize(file_path) / 1024  # KB
            
            info = f"檔案名稱: {os.path.basename(file_path)}\n"
            info += f"檔案大小: {file_size:.2f} KB\n"
            info += f"段落數量: {paragraphs}\n"
            info += f"表格數量: {tables}\n"
            info += f"圖片數量: {image_count}\n"
            info += f"章節數量: {sections}\n"
            
            self.file_info.setText(info)
        except Exception as e:
            self.file_info.setText(f"無法讀取檔案資訊: {str(e)}")
    
    def start_conversion(self):
        """開始轉換Word到PDF"""
        if not self.input_file or not self.output_file:
            QMessageBox.warning(self, '警告', '請選擇輸入和輸出檔案')
            return
        
        # 禁用按鈕，避免重複操作
        self.convert_btn.setEnabled(False)
        self.browse_btn.setEnabled(False)
        self.save_btn.setEnabled(False)
        
        # 開始轉換執行緒
        self.conversion_thread = ConversionThread(self.input_file, self.output_file)
        self.conversion_thread.progress_signal.connect(self.update_progress)
        self.conversion_thread.finished_signal.connect(self.conversion_finished)
        self.conversion_thread.error_signal.connect(self.conversion_error)
        self.conversion_thread.start()
    
    def update_progress(self, value):
        """更新進度條"""
        self.progress_bar.setValue(value)
    
    def conversion_finished(self, output_file):
        """轉換完成處理"""
        self.progress_label.setText('轉換完成!')
        
        # 重新啟用按鈕
        self.convert_btn.setEnabled(True)
        self.browse_btn.setEnabled(True)
        self.save_btn.setEnabled(True)
        self.open_btn.setEnabled(True)
        
        QMessageBox.information(
            self, '完成', f'Word檔案已成功轉換為PDF!\n儲存於: {output_file}'
        )
    
    def conversion_error(self, error_msg):
        """轉換錯誤處理"""
        self.progress_label.setText('轉換失敗!')
        self.progress_bar.setValue(0)
        
        # 重新啟用按鈕
        self.convert_btn.setEnabled(True)
        self.browse_btn.setEnabled(True)
        self.save_btn.setEnabled(True)
        
        QMessageBox.critical(self, '錯誤', f'轉換過程中發生錯誤:\n{error_msg}')
    
    def open_output_file(self):
        """開啟生成的PDF檔案"""
        if os.path.exists(self.output_file):
            QDesktopServices.openUrl(QUrl.fromLocalFile(self.output_file))
        else:
            QMessageBox.warning(self, '警告', '找不到輸出檔案')


if __name__ == '__main__':
    app = QApplication(sys.argv)
    converter = WordToPdfConverter()
    converter.show()
    sys.exit(app.exec_())
